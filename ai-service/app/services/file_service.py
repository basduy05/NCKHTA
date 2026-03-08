"""
File text extraction service.
Supports: TXT, PDF (PyPDF2), DOCX (python-docx).
Falls back gracefully if libraries are not installed.
"""
import os
import io
import traceback


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text content from uploaded file.
    Supports: .txt, .pdf, .docx, .doc, .md, .csv
    Returns extracted text as a string.
    """
    ext = os.path.splitext(filename)[1].lower()
    
    if ext in (".txt", ".md", ".csv", ".log"):
        return _extract_text(file_bytes)
    elif ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)
    else:
        # Try as plain text
        return _extract_text(file_bytes)


def _extract_text(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text.strip())
        result = "\n\n".join(texts)
        if result.strip():
            return result
        return "(PDF has no extractable text - may be scanned/image-based)"
    except ImportError:
        return "(PyPDF2 not installed - cannot extract PDF. Install: pip install PyPDF2)"
    except Exception as e:
        print(f"[FILE SERVICE] PDF extraction error: {e}")
        traceback.print_exc()
        return f"(PDF extraction failed: {str(e)})"


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs) if paragraphs else "(Empty DOCX document)"
    except ImportError:
        return "(python-docx not installed - cannot extract DOCX. Install: pip install python-docx)"
    except Exception as e:
        print(f"[FILE SERVICE] DOCX extraction error: {e}")
        traceback.print_exc()
        return f"(DOCX extraction failed: {str(e)})"


def get_supported_extensions() -> list:
    """Return list of supported file extensions."""
    return [".txt", ".md", ".csv", ".pdf", ".docx", ".doc", ".log"]
